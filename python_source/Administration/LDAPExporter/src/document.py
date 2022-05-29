from docx import Document
from docx.shared import Mm
from docx.oxml.ns import qn
from docx.shared import Pt


class WordDocument:
    def __init__(self):
        margin = 20
        self.document = Document()
        section = self.document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(margin)
        section.right_margin = Mm(margin)
        section.top_margin = Mm(margin)
        section.bottom_margin = Mm(margin)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        # setting the page layout to have 3 columns
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')

    def create_document(self):
        return self.document

    def add_paragraph_text(self, paragraph, text, is_bold=False, font_size=12):
        run = paragraph.add_run(text)
        font = run.font
        font.size = Pt(font_size)
        if is_bold:
            font.bold = True

    def add_paragraph(self, document):
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        return paragraph

    def save_document(self, document, output_path):
        document.save(output_path)