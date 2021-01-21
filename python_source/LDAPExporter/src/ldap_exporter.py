import os
from docx import Document
from docx.shared import Mm
from docx.oxml.ns import qn
from src.util.constants import Constants


class LdapExporter:
    def __init__(self):
        self.file_name = ""

    def convert_ldap_data(self, data_file_path):
        # get the filename only
        self.file_name = os.path.basename(data_file_path)
        data_dict = {}
        attrs = ["facsimileTelephoneNumber",
                 "st",
                 "street",
                 "email",
                 "mail",
                 "mobile",
                 "registeredAddress",
                 "telephoneNumber"]

        with open(data_file_path, "r") as data:
            data_lines = data.readlines()
            levels_split_cleaned = ""
            for line in data_lines:
                if line.startswith("dn"):
                    levels = line.split(":")[1].strip()
                    levels_split = levels.split(",")
                    levels_split_cleaned = list(
                        map(lambda x: x.replace("ou=", "").replace("dc=", "").replace("cn=", "").replace("o=", ""),
                            levels_split))

                    # this will create a dictionary from a list. it checks if the combination of keys
                    # in the list already exists inside the nested dictionary. if not it adds it
                    # inside the nested dictionary.
                    _data_dict = data_dict
                    for level in reversed(levels_split_cleaned):
                        try:
                            _data_dict = _data_dict[level]
                        except KeyError:
                            _data_dict[level] = {}
                            _data_dict = _data_dict[level]

                else:
                    for attr in attrs:
                        if line.startswith(attr):
                            value = line.split(attr + ":")[1]
                            key = attr.strip()
                            value = value.strip()

                            _data_dict = data_dict
                            for level in reversed(levels_split_cleaned):
                                _data_dict = _data_dict[level]
                            if key in _data_dict:
                                if not isinstance(_data_dict[key], list):
                                    _data_dict[key] = [_data_dict[key]]
                                _data_dict[key].append(value)
                            else:
                                _data_dict[key] = value
            return data_dict

    def print_nested_dict(self, d):
        # prints all the keys and values of a dictionary
        for k, v in d.items():
            if isinstance(v, dict):
                yield k
                yield from self.print_nested_dict(v)
            else:
                yield k, v

    def export_data(self, data_dict):
        output_path = os.path.join(Constants.OUTPUT_DIR, os.path.splitext(self.file_name)[0] + ".docx")
        root_data_dict = data_dict['design']['ptcc']
        doc = self.create_document()
        par = self.add_paragraph(doc)
        for entry in self.print_nested_dict(root_data_dict):
            if isinstance(entry, tuple):
                key, value = entry
                if isinstance(value, list):
                    value = "/\n".join(value)
                    text = "{}:\n{}\n".format(key, value)
                else:
                    text = "{}: {}\n".format(key, value)
                bold = False
            else:
                text = "\n" + entry + "\n"
                bold = True
            self.add_paragraph_text(par, text, bold)
        self.save_document(doc, output_path)

    # document related methods (SHOULD have been a separate class)
    def create_document(self):
        margin = 20
        document = Document()
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(margin)
        section.right_margin = Mm(margin)
        section.top_margin = Mm(margin)
        section.bottom_margin = Mm(margin)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        # setting the page layout to have 3 columns
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '2')
        return document

    def add_paragraph_text(self, paragraph, text, is_bold=False):
        if is_bold:
            paragraph.add_run(text).bold = True
        else:
            paragraph.add_run(text)

    def add_paragraph(self, document):
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']
        return paragraph

    def save_document(self, document, output_path):
        document.save(output_path)
