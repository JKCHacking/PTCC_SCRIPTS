import os
import csv
from docx import Document
from docx.shared import Mm
from docx.oxml.ns import qn
from src.util.constants import Constants
from src.util.watcher import Watcher


class LDAPGenerator:
    def __init__(self):
        pass

    def iter_input(self):
        input_dir = Constants.INPUT_DIR
        for dir_path, dir_names, file_names in os.walk(input_dir):
            for file_name in file_names:
                file_full_path = os.path.join(dir_path, file_name)
                if file_full_path.endswith(Constants.CSV_FILE_EXT):
                    data = self.get_csv_data(file_full_path)
                    # sort data according to first element
                    data.sort()
                    filename_ext = os.path.basename(file_full_path)
                    filename = os.path.splitext(filename_ext)[0]
                    output_dir = os.path.join(Constants.OUTPUT_DIR, filename + Constants.DOCX_FILE_EXT)
                    self.create_document(data, output_dir)

    def get_csv_data(self, file_full_path):
        data = []
        with open(file_full_path, newline='') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                if "cn=" in row[0]:
                    row_data = []
                    # get the name of the contact
                    try:
                        contact_name = row[0].split(',')[0].lstrip("cn=")
                        row_data.append(contact_name)
                        street = row[1]
                        row_data.append(street)
                        city = row[2]
                        row_data.append(city)
                        country = row[3]
                        row_data.append(country)
                        fax = row[4]
                        row_data.append(fax)
                        tel_num = row[5]
                        row_data.append(tel_num)
                        mobile = row[6]
                        row_data.append(mobile)
                        email = row[7]
                        row_data.append(email)
                    except IndexError:
                        pass
                    data.append(row_data)
        return data

    def create_document(self, data, output_path):
        margin = 20
        document = Document()
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(margin)
        section.right_margin = Mm(margin)
        section.top_margin = Mm(margin)
        section.bottom_margin = Mm(margin)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Normal']

        # setting the page layout to have 3 columns
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '3')

        # this represent the starting letter of each names.
        current_char = ''
        for row in data:
            contact_name = row[0]
            street = row[1]
            city = row[2]
            country = row[3]
            fax = row[4]
            tel_num = row[5]
            mobile = row[6]
            email = row[7]

            watcher = Watcher(current_char)
            current_char = contact_name[0]
            watcher.set_value(current_char)
            if watcher.has_changed():
                if current_char.isnumeric():
                    current_char = "123"

                paragraph.add_run(current_char + "\n\n").bold = True

            if contact_name:
                paragraph.add_run(contact_name + "\n").bold = True
            if street:
                paragraph.add_run(street + "\n")
            if city:
                paragraph.add_run(city + "\n")
            if country:
                paragraph.add_run(country + "\n")
            if fax:
                fax_list = fax.split("|")
                fax_text = "/\n".join(fax_list)
                paragraph.add_run("Fax: " + fax_text + "\n")
            if tel_num:
                tel_num_list = tel_num.split("|")
                tel_num_text = "/\n".join(tel_num_list)
                paragraph.add_run("Telephone: " + tel_num_text + "\n")
            if mobile:
                mobile_list = mobile.split("|")
                mobile_text = "/\n".join(mobile_list)
                paragraph.add_run("Mobile: " + mobile_text + "\n")
            if email:
                email_list = email.split("|")
                email_text = "/\n".join(email_list)
                paragraph.add_run("E-Mail: " + email_text + "\n")
            paragraph.add_run("\n")
        document.save(output_path)
